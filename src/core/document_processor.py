import os
from typing import List, Dict, Optional, Callable
from pathlib import Path
import PyPDF2
from docx import Document
from datetime import datetime

class DocumentProcessor:
    def __init__(self):
        self.supported_extensions = {
            '.txt': self._process_txt,
            '.pdf': self._process_pdf,
            '.doc': self._process_doc,
            '.docx': self._process_docx
        }

    def process_document(self, file_path: str, progress_callback: Optional[Callable[[int], None]] = None) -> List[Dict]:
        """Process single document"""
        try:
            path = Path(file_path)
            if not path.exists():
                raise FileNotFoundError(f"File does not exist: {file_path}")

            if path.suffix.lower() not in self.supported_extensions:
                raise ValueError(f"Unsupported file type: {path.suffix}")

            # Call corresponding processing function
            processor = self.supported_extensions[path.suffix.lower()]
            chunks = processor(file_path, progress_callback)

            # Add metadata
            for chunk in chunks:
                chunk.update({
                    "source": file_path,
                    "filename": path.name,
                    "file_type": path.suffix[1:].upper(),
                    "created_at": datetime.now().isoformat()
                })

            return chunks

        except Exception as e:
            raise Exception(f"Failed to process document: {str(e)}")

    def _process_txt(self, file_path: str, progress_callback: Optional[Callable[[int], None]] = None) -> List[Dict]:
        """Process TXT file"""
        chunks = []
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()

            # Simple paragraph processing
            paragraphs = content.split('\n\n')
            total = len(paragraphs)

            for i, para in enumerate(paragraphs, 1):
                if para.strip():
                    chunks.append({
                        "content": para.strip(),
                        "chunk_type": "paragraph",
                        "chunk_index": i
                    })

                if progress_callback:
                    progress_callback(int(i * 100 / total))

            return chunks

        except UnicodeDecodeError:
            # Try other encodings
            encodings = ['gbk', 'gb2312', 'utf-16', 'ascii']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    
                    paragraphs = content.split('\n\n')
                    total = len(paragraphs)

                    for i, para in enumerate(paragraphs, 1):
                        if para.strip():
                            chunks.append({
                                "content": para.strip(),
                                "chunk_type": "paragraph",
                                "chunk_index": i
                            })

                        if progress_callback:
                            progress_callback(int(i * 100 / total))

                    return chunks
                except UnicodeDecodeError:
                    continue

            raise Exception(f"Unable to decode file: {file_path}")

    def _process_pdf(self, file_path: str, progress_callback: Optional[Callable[[int], None]] = None) -> List[Dict]:
        """Process PDF file"""
        chunks = []
        try:
            with open(file_path, 'rb') as f:
                pdf = PyPDF2.PdfReader(f)
                total_pages = len(pdf.pages)

                for i, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text.strip():
                        chunks.append({
                            "content": text.strip(),
                            "chunk_type": "page",
                            "chunk_index": i,
                            "page_number": i
                        })

                    if progress_callback:
                        progress_callback(int(i * 100 / total_pages))

            return chunks

        except Exception as e:
            raise Exception(f"Failed to process PDF file: {str(e)}")

    def _process_docx(self, file_path: str, progress_callback: Optional[Callable[[int], None]] = None) -> List[Dict]:
        """Process DOCX file"""
        chunks = []
        try:
            doc = Document(file_path)
            total_paras = len(doc.paragraphs)

            for i, para in enumerate(doc.paragraphs, 1):
                if para.text.strip():
                    chunks.append({
                        "content": para.text.strip(),
                        "chunk_type": "paragraph",
                        "chunk_index": i,
                        "style": para.style.name
                    })

                if progress_callback:
                    progress_callback(int(i * 100 / total_paras))

            return chunks

        except Exception as e:
            raise Exception(f"Failed to process DOCX file: {str(e)}")

    def _process_doc(self, file_path: str, progress_callback: Optional[Callable[[int], None]] = None) -> List[Dict]:
        """Process DOC file"""
        # Currently not supported directly, recommend converting to DOCX
        raise NotImplementedError("Currently not supported directly processing DOC file, please convert to DOCX format") 